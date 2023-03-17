import pandas as pd
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor


def import_excel(file_path):
    client_export = pd.read_excel(file_path)

    client_export = \
        client_export[(client_export['Status'] == "Actueel") & (client_export['Opgenomen in Specificatie?'] == "Ja")]

    return client_export


def tabel_maken(df):
    doc = Document('word_file.docx')
    for index, row in df.iterrows():
        table = doc.add_table(rows=5, cols=2)
        table.style = 'PythonStyle'
        m_cell_1 = table.cell(1, 0)
        m_cell_2 = table.cell(1, 1)
        merge_1 = m_cell_1.merge(m_cell_2)
        m_cell_3 = table.cell(2, 0)
        m_cell_4 = table.cell(2, 1)
        merge_2 = m_cell_3.merge(m_cell_4)
        m_cell_5 = table.cell(3, 0)
        m_cell_6 = table.cell(3, 1)
        merge_3 = m_cell_5.merge(m_cell_6)
        m_cell_7 = table.cell(4, 0)
        m_cell_8 = table.cell(4, 1)
        merge_4 = m_cell_7.merge(m_cell_8)

        # passing variable to non merged cells
        cell_1 = table.cell(0, 0)
        cell_2 = table.cell(0, 1)

        doc.add_paragraph()

        # Populate the table with data from the current row
        for j, col in enumerate(df.columns):
            cell_value = row[col]

            # Paste the cell value into the appropriate table cell
            if j == 0:
                if cell_value != "":
                    cell_1.text = str('Eis-ID: ' + '\n' + str(cell_value))
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
                    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
                    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # White color
                    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                else:
                    cell_1.text = 'Eis-ID: '
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
                    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
                    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # White color
                    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True

            elif j == 1:
                if cell_value != "":
                    merge_1.text = str('Eistekst: ' + '\n' + str(cell_value))
                else:
                    merge_1.text = 'Eistekst: '
            elif j == 2:
                if cell_value != "":
                    merge_2.text = str(
                        'Eistoelichting: ' + '\n' + str(cell_value))
                else:
                    merge_2.text = 'Eistoelichting: '
            elif j == 3:
                if cell_value != "":
                    cell_2.text = str('Onderwerp: ' + str(cell_value))
                else:
                    cell_2.text = 'Onderwerp: '
            elif j == 4:
                if cell_value == "VOLDOET" or "Voldoet" or "voldoet" or "VOLDOET NIET" or "Voldoet niet" or \
                        "voldoet niet":
                    merge_3.text = f'Oordeel verificatie: {str(cell_value)}'

                else:
                    merge_3.text = 'Oordeel verificatie:'
            elif j == 5:
                if cell_value != "":
                    merge_4.text = str(
                        'Toelichting oordeel verificatie:' + '\n' + str(cell_value))
                else:
                    merge_4.text = 'Toelichting oordeel verificatie:'

    return doc
